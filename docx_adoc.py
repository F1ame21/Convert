import docx
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
import copy
import numpy as np

def data_table(doc):
  all_tables = doc.tables
  data_tables = {i:None for i in range(len(all_tables))}
  for i, table in enumerate(all_tables):
      data_tables[i] = [[] for _ in range(len(table.rows))]
      for j, row in enumerate(table.rows):
          for cell in row.cells:
              data_tables[i][j].append(cell.text)
  return data_tables

def iter_block_items(parent):
    if isinstance(parent, docx.document.Document): 
        parent_elm = parent.element.body
    else:
        raise ValueError("something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield docx.text.paragraph.Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield docx.table.Table(child, parent)

def check_Title_style(k, doc):
  text = doc.paragraphs[k].text
  text = '= ' + text
  return text

def check_Heading_style(k, doc, style):
  text = doc.paragraphs[k].text
  s = int(style[-1]) + 1
  text = (s * '=') + ' ' + text
  return text

def check_style_text(k, doc):
  text = copy.deepcopy(doc.paragraphs[k])
  for run in (text.runs):
    xmlstr = str(run.element.xml)
    if 'pic:pic' in run.element.xml:
      run.text = 'Photo'
      break  
  return text.text

def check_List_Paragraph_style(k, doc):
  list_item = copy.deepcopy(doc.paragraphs[k])
  list_lvl = doc.paragraphs[k]._p.pPr.numPr.ilvl.val
  list_item = ((list_lvl + 1) * '*') + ' ' + str(list_item.text)
  return list_item

def append_table(table, total):
  line = '|==='
  total.append(line)
  np_table = np.array(table)
  for i in range(np.shape(np_table)[0]):
    line = ''
    for j in range(np.shape(np_table)[1]):
      if np_table[i,j] == '':
        line = line + '|' + '  '
      else:
        line = line + '|'+ np_table[i,j] + ' '
    total.append(line)
  line = '|==='
  total.append(line)  
  return total

def write_in_asccidoc_file(d):
  file = open("convert.adoc", 'a+')
  data = iter(d)
  for item in data:
    if item == '|===':
      file.write("%s\n" % item)
      next_item = next(data)
      while next_item != '|===':
        file.write("%s\n" % next_item)
        next_item = next(data)
      file.write("%s\n\n" % item)
    else:
      file.write("%s\n\n" % item)
  file.close()
  return file

def DocxToAdoc(doc, total = [], k=0, number_table = 0):  
  data_tables = data_table(doc)
  for block in iter_block_items(doc):
    if 'text' in str(block):
      style = str(block.style.name)
      if 'Title' in style:
        title = check_Title_style(k, doc)
        total.append(title)
      elif "Heading" in style:
        heading = check_Heading_style(k, doc, style)
        total.append(heading)
      elif 'Normal' in style:
        txt = check_style_text(k, doc)
        if txt == '':
          pass
        else:
          total.append(txt)
      elif 'List Paragraph' in style:
          List = check_List_Paragraph_style(k, doc)
          total.append(List)
      k += 1
    elif 'table' in str(block):
      table = data_tables[number_table]
      total = append_table(table, total)
      number_table += 1
    else:
        print('have unknown value')
  write_in_asccidoc_file(total)
  return total
total = []
doc = Document('3.docx')
data = []
data = DocxToAdoc(doc)